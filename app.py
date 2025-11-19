from flask import Flask, render_template_string, request, send_file
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml import OxmlElement, ns
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
import io
import os
import json

app = Flask(__name__)

# HTML avec interface améliorée
HTML = """
<!DOCTYPE html>
<html lang="ar" dir="rtl">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>إنشاء جدول DOCX</title>
    <style>
        body { 
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; 
            text-align: right; 
            padding: 20px; 
            background-color: #f5f5f5;
            margin: 0;
        }
        .container {
            max-width: 1200px;
            margin: 0 auto;
            background: white;
            padding: 30px;
            border-radius: 10px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        }
        h2 {
            color: #2c3e50;
            text-align: center;
            margin-bottom: 30px;
        }
        .form-group {
            margin-bottom: 20px;
        }
        label {
            display: block;
            margin-bottom: 8px;
            font-weight: bold;
            color: #34495e;
        }
        input, select {
            width: 100%;
            margin: 5px 0;
            padding: 12px;
            border: 1px solid #ddd;
            border-radius: 5px;
            font-size: 16px;
        }
        .criteria-section {
            display: flex;
            gap: 20px;
            margin-top: 20px;
        }
        .suggested-criteria, .selected-criteria {
            flex: 1;
        }
        .suggested-container, .selected-container {
            border: 2px dashed #3498db;
            padding: 15px;
            border-radius: 5px;
            margin: 10px 0;
            min-height: 200px;
            background-color: #f8f9fa;
        }
        .suggested-item {
            background: #2ecc71;
            color: white;
            padding: 10px 15px;
            margin: 5px;
            border-radius: 20px;
            display: block;
            cursor: move;
            text-align: center;
        }
        .suggested-item:hover {
            background: #27ae60;
        }
        .criteria-item {
            background: #3498db;
            color: white;
            padding: 10px 15px;
            margin: 5px;
            border-radius: 20px;
            display: flex;
            justify-content: space-between;
            align-items: center;
            cursor: move;
        }
        .criteria-item:hover {
            background: #2980b9;
        }
        .criteria-actions {
            display: flex;
            gap: 5px;
        }
        .action-btn {
            background: rgba(255,255,255,0.2);
            border: none;
            color: white;
            padding: 5px 8px;
            border-radius: 50%;
            cursor: pointer;
            font-size: 12px;
        }
        .action-btn:hover {
            background: rgba(255,255,255,0.3);
        }
        .btn {
            background: #3498db;
            color: white;
            border: none;
            padding: 15px 30px;
            border-radius: 5px;
            font-size: 18px;
            cursor: pointer;
            width: 100%;
            margin-top: 20px;
        }
        .btn:hover {
            background: #2980b9;
        }
        .btn-secondary {
            background: #95a5a6;
            padding: 10px 20px;
            font-size: 14px;
            width: auto;
        }
        .btn-secondary:hover {
            background: #7f8c8d;
        }
        .btn-danger {
            background: #e74c3c;
            padding: 8px 15px;
            font-size: 12px;
            width: auto;
        }
        .btn-danger:hover {
            background: #c0392b;
        }
        .drag-info {
            text-align: center;
            color: #7f8c8d;
            font-style: italic;
            margin: 10px 0;
        }
        .instructions {
            background: #fff3cd;
            border: 1px solid #ffeaa7;
            padding: 10px;
            border-radius: 5px;
            margin: 10px 0;
            font-size: 14px;
        }
        .section-title {
            background: #34495e;
            color: white;
            padding: 10px;
            border-radius: 5px;
            text-align: center;
            margin-bottom: 10px;
        }
        .empty-message {
            text-align: center;
            color: #7f8c8d;
            font-style: italic;
            padding: 20px;
        }
        .table-preview {
            margin-top: 20px;
            border: 2px solid #3498db;
            border-radius: 5px;
            padding: 15px;
            background: white;
        }
        .preview-table {
            width: 100%;
            border-collapse: collapse;
            font-size: 12px;
        }
        .preview-table th, .preview-table td {
            border: 1px solid #ddd;
            padding: 5px;
            text-align: center;
        }
        .preview-table th {
            background-color: #f8f9fa;
            font-weight: bold;
        }
        .option-group {
            background: #e8f6f3;
            padding: 15px;
            border-radius: 5px;
            border: 1px solid #27ae60;
            margin: 10px 0;
        }
        .checkbox-group {
            display: flex;
            align-items: center;
            gap: 10px;
            margin: 5px 0;
        }
        .edit-form {
            background: #f8f9fa;
            padding: 10px;
            border-radius: 5px;
            margin: 5px 0;
            border: 1px dashed #3498db;
        }
        .edit-input {
            width: 100%;
            padding: 8px;
            border: 1px solid #ddd;
            border-radius: 3px;
            margin-bottom: 5px;
        }
        .indicator-option {
            background: #fff3cd;
            padding: 8px;
            border-radius: 3px;
            margin: 5px 0;
            border: 1px solid #ffeaa7;
        }
        .hidden {
            display: none;
        }
        .other-subject-input {
            margin-top: 10px;
            padding: 10px;
            background: #f8f9fa;
            border-radius: 5px;
            border: 1px solid #3498db;
        }
        .loading {
            display: none;
            text-align: center;
            margin: 20px 0;
        }
        .spinner {
            border: 4px solid #f3f3f3;
            border-top: 4px solid #3498db;
            border-radius: 50%;
            width: 40px;
            height: 40px;
            animation: spin 2s linear infinite;
            margin: 0 auto;
        }
        @keyframes spin {
            0% { transform: rotate(0deg); }
            100% { transform: rotate(360deg); }
        }
    </style>
</head>
<body>
    <div class="container">
        <h2>إنشاء جدول التقييم DOCX</h2>
        <form method="POST" id="docxForm">
            <div class="form-group">
                <label>القسم:</label>
                <input type="text" name="classe" value="سنة رابعة" required placeholder="أدخل اسم القسم">
            </div>
            
            <div class="form-group">
                <label>نوع التقييم:</label>
                <select id="matiere" name="matiere" required onchange="handleSubjectChange()">
                    <option value="">اختر نوع التقييم</option>
                    <option value="التواصل الشفوي">التواصل الشفوي</option>
                    <option value="القراءة">القراءة</option>
                    <option value="الإنتاج الكتابي">الإنتاج الكتابي</option>
                    <option value="قواعد اللغة">قواعد اللغة</option>
                    <option value="أخرى">أخرى</option>
                </select>
                
                <div id="otherSubjectInput" class="other-subject-input hidden">
                    <label>اسم المادة:</label>
                    <input type="text" id="otherSubjectName" name="other_subject_name" placeholder="أدخل اسم المادة">
                </div>
            </div>
            
            <div class="form-group">
                <label>إعداد المعايير:</label>
                
                <div class="instructions">
                    💡 <strong>تعليمات:</strong> 
                    <br>• اختر نوع التقييم أولاً
                    <br>• اسحب المعايير من القائمة المقترحة إلى قائمة المعايير المختارة
                    <br>• انقر على ✏️ لتعديل اسم المعيار
                    <br>• انقر على 📊 لاختيار المؤشرات لكل معيار
                </div>

                <div class="criteria-section">
                    <!-- القائمة المقترحة -->
                    <div class="suggested-criteria">
                        <div class="section-title">المعايير المقترحة</div>
                        <div class="suggested-container" id="suggestedContainer" ondragover="allowDrop(event)" ondrop="dropInSuggested(event)">
                            <div class="drag-info">اسحب المعايير إلى القائمة المختارة</div>
                            <div id="suggestedList"></div>
                        </div>
                    </div>
                    
                    <!-- القائمة المختارة -->
                    <div class="selected-criteria">
                        <div class="section-title">المعايير المختارة</div>
                        <div class="selected-container" id="selectedContainer" ondragover="allowDrop(event)" ondrop="dropInSelected(event)">
                            <div class="drag-info">اسحب المعايير هنا</div>
                            <div id="selectedList"></div>
                        </div>
                    </div>
                </div>
                
                <input type="hidden" name="criteria" id="criteriaInput" required>
                <input type="hidden" name="indicators_config" id="indicatorsConfigInput" value="{}">
                
                <div class="criteria-actions" style="justify-content: center; margin-top: 20px;">
                    <button type="button" class="btn-danger" onclick="clearAllCriteria()">حذف الكل</button>
                </div>
            </div>

            <!-- معاينة الجدول -->
            <div class="form-group">
                <div class="table-preview">
                    <div class="section-title">معاينة الجدول</div>
                    <div id="tablePreview">
                        <div class="empty-message">سيظهر معاينة الجدول هنا بعد اختيار المعايير</div>
                    </div>
                </div>
            </div>
            
            <div class="form-group">
                <label>اختر مجموعة التلاميذ:</label>
                <select name="group_choice" required>
                    <option value="1">المجموعة السابقة</option>
                    <option value="2">المجموعة الجديدة</option>
                </select>
            </div>
            
            <div class="loading" id="loadingIndicator">
                <div class="spinner"></div>
                <p>جاري إنشاء الملف، يرجى الانتظار...</p>
            </div>
            
            <button type="submit" class="btn" id="submitBtn">إنشاء الملف</button>
        </form>
    </div>

    <script>
        let selectedCriteria = [];
        let suggestedCriteria = [];
        let editingIndex = -1;
        let indicatorsConfig = {}; // {criteriaName: {useIndicators: boolean, indicatorNames: []}}
        
        const subjectCriteria = {
            "التواصل الشفوي": [
                "الملائمة", "التغنيم", "الانسجام", "الاتساق", "الثراء"
            ],
            "القراءة": [
                "القراءة الجهرية", "معالجة النص", "التصرف في النص", "إبداء الرأي"
            ],
            "الإنتاج الكتابي": [
                "الملائمة", "سلامة بناء النص", "المقروئية", "ثراء اللغة والطرافة"
            ],
            "قواعد اللغة": [
                "التعرف على الظاهرة اللغوية", "توظيف الظاهرة اللغوية"
            ],
            "أخرى": [
                "معيار 1", "معيار 2", "معيار 3"
            ]
        };
        
        function handleSubjectChange() {
            const subject = document.getElementById('matiere').value;
            const otherInput = document.getElementById('otherSubjectInput');
            
            if (subject === 'أخرى') {
                otherInput.classList.remove('hidden');
            } else {
                otherInput.classList.add('hidden');
            }
            
            updateSuggestedCriteria();
        }
        
        function updateCriteriaInput() {
            document.getElementById('criteriaInput').value = JSON.stringify(selectedCriteria);
            document.getElementById('indicatorsConfigInput').value = JSON.stringify(indicatorsConfig);
            updateTablePreview();
        }
        
        function updateSuggestedCriteria() {
            const subject = document.getElementById('matiere').value;
            const suggestedList = document.getElementById('suggestedList');
            
            suggestedCriteria = subjectCriteria[subject] || [];
            suggestedList.innerHTML = '';
            
            if (suggestedCriteria.length === 0) {
                suggestedList.innerHTML = '<div class="empty-message">لا توجد معايير مقترحة</div>';
                return;
            }
            
            suggestedCriteria.forEach(criteria => {
                if (!selectedCriteria.includes(criteria)) {
                    const item = document.createElement('div');
                    item.className = 'suggested-item';
                    item.textContent = criteria;
                    item.draggable = true;
                    item.ondragstart = (e) => dragStart(e, criteria, 'suggested');
                    suggestedList.appendChild(item);
                }
            });
            
            if (suggestedList.children.length === 0) {
                suggestedList.innerHTML = '<div class="empty-message">جميع المعايير مضافة</div>';
            }
        }
        
        function addToSelected(criteria) {
            if (!selectedCriteria.includes(criteria)) {
                selectedCriteria.push(criteria);
                // إعداد افتراضي للمؤشرات
                if (!indicatorsConfig[criteria]) {
                    indicatorsConfig[criteria] = {
                        useIndicators: false,
                        indicatorNames: ["مؤشر 1", "مؤشر 2", "مؤشر 3"]
                    };
                }
                renderSelectedCriteria();
                updateSuggestedCriteria();
            }
        }
        
        function removeFromSelected(index) {
            const criteria = selectedCriteria[index];
            delete indicatorsConfig[criteria];
            selectedCriteria.splice(index, 1);
            renderSelectedCriteria();
            updateSuggestedCriteria();
        }
        
        function startEdit(index) {
            editingIndex = index;
            renderSelectedCriteria();
        }
        
        function saveEdit(index, newValue) {
            const oldValue = selectedCriteria[index];
            if (newValue.trim() && !selectedCriteria.includes(newValue.trim())) {
                // تحديث التكوين إذا تغير الاسم
                if (indicatorsConfig[oldValue]) {
                    indicatorsConfig[newValue.trim()] = indicatorsConfig[oldValue];
                    delete indicatorsConfig[oldValue];
                }
                selectedCriteria[index] = newValue.trim();
            }
            editingIndex = -1;
            renderSelectedCriteria();
            updateSuggestedCriteria();
        }
        
        function cancelEdit() {
            editingIndex = -1;
            renderSelectedCriteria();
        }
        
        function toggleIndicators(criteria) {
            if (!indicatorsConfig[criteria]) {
                indicatorsConfig[criteria] = {
                    useIndicators: true,
                    indicatorNames: ["مؤشر 1", "مؤشر 2", "مؤشر 3"]
                };
            } else {
                indicatorsConfig[criteria].useIndicators = !indicatorsConfig[criteria].useIndicators;
            }
            updateCriteriaInput();
        }
        
        function editIndicatorNames(criteria) {
            if (!indicatorsConfig[criteria]) {
                indicatorsConfig[criteria] = {
                    useIndicators: true,
                    indicatorNames: ["مؤشر 1", "مؤشر 2", "مؤشر 3"]
                };
            }
            
            const newNames = [];
            for (let i = 0; i < 3; i++) {
                const currentName = indicatorsConfig[criteria].indicatorNames[i] || `مؤشر ${i+1}`;
                const newName = prompt(`أدخل اسم المؤشر ${i+1} لـ "${criteria}":`, currentName);
                if (newName === null) return; // المستخدم ألغى
                newNames.push(newName.trim() || `مؤشر ${i+1}`);
            }
            
            indicatorsConfig[criteria].indicatorNames = newNames;
            updateCriteriaInput();
        }
        
        function renderSelectedCriteria() {
            const selectedList = document.getElementById('selectedList');
            selectedList.innerHTML = '';
            
            if (selectedCriteria.length === 0) {
                selectedList.innerHTML = '<div class="empty-message">لم يتم اختيار أي معايير</div>';
                updateCriteriaInput();
                return;
            }
            
            selectedCriteria.forEach((criteria, index) => {
                if (editingIndex === index) {
                    // وضع التعديل
                    const editForm = document.createElement('div');
                    editForm.className = 'edit-form';
                    editForm.innerHTML = `
                        <input type="text" 
                               class="edit-input" 
                               value="${criteria}" 
                               id="editInput-${index}"
                               placeholder="أدخل اسم المعيار">
                        <div style="display: flex; gap: 5px; justify-content: center;">
                            <button type="button" class="btn-secondary" onclick="saveEdit(${index}, document.getElementById('editInput-${index}').value)">حفظ</button>
                            <button type="button" class="btn-danger" onclick="cancelEdit()">إلغاء</button>
                        </div>
                    `;
                    selectedList.appendChild(editForm);
                    
                    // تركيز على حقل الإدخال
                    setTimeout(() => {
                        const input = document.getElementById(`editInput-${index}`);
                        input.focus();
                        input.select();
                    }, 100);
                } else {
                    // عرض عادي
                    const item = document.createElement('div');
                    item.className = 'criteria-item';
                    item.draggable = true;
                    item.ondragstart = (e) => dragStart(e, criteria, 'selected');
                    
                    const criteriaText = document.createElement('span');
                    criteriaText.textContent = criteria;
                    
                    const actions = document.createElement('div');
                    actions.className = 'criteria-actions';
                    
                    // زر المؤشرات
                    const indicatorsBtn = document.createElement('button');
                    indicatorsBtn.className = 'action-btn';
                    indicatorsBtn.type = 'button'; // لمنع الإرسال التلقائي
                    indicatorsBtn.innerHTML = indicatorsConfig[criteria]?.useIndicators ? '📊✅' : '📊';
                    indicatorsBtn.title = indicatorsConfig[criteria]?.useIndicators ? 'المؤشرات مفعلة - انقر لإلغاء' : 'إضافة مؤشرات';
                    indicatorsBtn.onclick = () => toggleIndicators(criteria);
                    actions.appendChild(indicatorsBtn);
                    
                    // زر تعديل أسماء المؤشرات
                    if (indicatorsConfig[criteria]?.useIndicators) {
                        const editIndicatorsBtn = document.createElement('button');
                        editIndicatorsBtn.className = 'action-btn';
                        editIndicatorsBtn.type = 'button'; // لمنع الإرسال التلقائي
                        editIndicatorsBtn.innerHTML = '✏️';
                        editIndicatorsBtn.title = 'تعديل أسماء المؤشرات';
                        editIndicatorsBtn.onclick = () => editIndicatorNames(criteria);
                        actions.appendChild(editIndicatorsBtn);
                    }
                    
                    // زر تعديل اسم المعيار
                    const editBtn = document.createElement('button');
                    editBtn.className = 'action-btn';
                    editBtn.type = 'button'; // لمنع الإرسال التلقائي
                    editBtn.innerHTML = '✏️';
                    editBtn.title = 'تعديل اسم المعيار';
                    editBtn.onclick = () => startEdit(index);
                    actions.appendChild(editBtn);
                    
                    // زر الحذف
                    const deleteBtn = document.createElement('button');
                    deleteBtn.className = 'action-btn';
                    deleteBtn.type = 'button'; // لمنع الإرسال التلقائي
                    deleteBtn.innerHTML = '🗑️';
                    deleteBtn.title = 'حذف';
                    deleteBtn.onclick = () => removeFromSelected(index);
                    actions.appendChild(deleteBtn);
                    
                    item.appendChild(criteriaText);
                    item.appendChild(actions);
                    selectedList.appendChild(item);
                    
                    // عرض حالة المؤشرات
                    if (indicatorsConfig[criteria]?.useIndicators) {
                        const indicatorInfo = document.createElement('div');
                        indicatorInfo.className = 'indicator-option';
                        indicatorInfo.innerHTML = `
                            <small>المؤشرات: ${indicatorsConfig[criteria].indicatorNames.join('، ')}</small>
                        `;
                        selectedList.appendChild(indicatorInfo);
                    }
                }
            });
            
            updateCriteriaInput();
        }
        
        function updateTablePreview() {
            const preview = document.getElementById('tablePreview');
            
            if (selectedCriteria.length === 0) {
                preview.innerHTML = '<div class="empty-message">سيظهر معاينة الجدول هنا بعد اختيار المعايير</div>';
                return;
            }
            
            let html = '<table class="preview-table">';
            
            // حساب عدد الأعمدة
            let totalCols = 1; // عمود الأسماء
            selectedCriteria.forEach(criteria => {
                if (indicatorsConfig[criteria]?.useIndicators) {
                    totalCols += 3; // 3 مؤشرات
                } else {
                    totalCols += 1; // معيار واحد بدون مؤشرات
                }
            });
            
            // بناء رأس الجدول
            html += '<tr>';
            html += '<th rowspan="2">اسم التلميذ</th>';
            
            selectedCriteria.forEach(criteria => {
                const useIndicators = indicatorsConfig[criteria]?.useIndicators;
                if (useIndicators) {
                    html += `<th colspan="3">${criteria}</th>`;
                } else {
                    html += `<th rowspan="2">${criteria}</th>`;
                }
            });
            html += '</tr>';
            
            // الصف الثاني للرأس (للمؤشرات فقط)
            const hasIndicators = selectedCriteria.some(criteria => indicatorsConfig[criteria]?.useIndicators);
            if (hasIndicators) {
                html += '<tr>';
                selectedCriteria.forEach(criteria => {
                    if (indicatorsConfig[criteria]?.useIndicators) {
                        const names = indicatorsConfig[criteria].indicatorNames;
                        html += `<th>${names[0]}</th><th>${names[1]}</th><th>${names[2]}</th>`;
                    }
                });
                html += '</tr>';
            }
            
            // صفوف التلاميذ (3 صفوف كمثال)
            for (let i = 1; i <= 3; i++) {
                html += '<tr>';
                html += `<td>التلميذ ${i}</td>`;
                selectedCriteria.forEach(criteria => {
                    if (indicatorsConfig[criteria]?.useIndicators) {
                        html += '<td></td><td></td><td></td>';
                    } else {
                        html += '<td></td>';
                    }
                });
                html += '</tr>';
            }
            
            html += '</table>';
            html += '<div style="text-align: center; margin-top: 10px; color: #7f8c8d; font-size: 12px;">';
            html += '🔹 = معيار بدون مؤشرات | 📊✅ = معيار مع مؤشرات';
            html += '</div>';
            
            preview.innerHTML = html;
        }
        
        function clearAllCriteria() {
            if (confirm('هل أنت متأكد من حذف جميع المعايير المختارة؟')) {
                selectedCriteria = [];
                indicatorsConfig = {};
                editingIndex = -1;
                renderSelectedCriteria();
                updateSuggestedCriteria();
            }
        }
        
        function allowDrop(ev) {
            ev.preventDefault();
        }
        
        function dragStart(ev, criteria, source) {
            ev.dataTransfer.setData("text/plain", JSON.stringify({
                criteria: criteria,
                source: source
            }));
        }
        
        function dropInSuggested(ev) {
            ev.preventDefault();
            const data = JSON.parse(ev.dataTransfer.getData("text/plain"));
            if (data.source === 'selected') {
                const index = selectedCriteria.indexOf(data.criteria);
                if (index > -1) {
                    removeFromSelected(index);
                }
            }
        }
        
        function dropInSelected(ev) {
            ev.preventDefault();
            const data = JSON.parse(ev.dataTransfer.getData("text/plain"));
            if (data.source === 'suggested') {
                addToSelected(data.criteria);
            }
        }
        
        // منع الإرسال التلقائي للنموذج
        document.getElementById('docxForm').addEventListener('submit', function(e) {
            e.preventDefault();
            
            // التحقق من صحة البيانات
            const subject = document.getElementById('matiere').value;
            if (subject === 'أخرى') {
                const otherSubject = document.getElementById('otherSubjectName').value.trim();
                if (!otherSubject) {
                    alert('يرجى إدخال اسم المادة');
                    return;
                }
            }
            
            if (selectedCriteria.length === 0) {
                alert('يرجى اختيار معايير التقييم');
                return;
            }
            
            // إظهار مؤشر التحميل
            const loadingIndicator = document.getElementById('loadingIndicator');
            const submitBtn = document.getElementById('submitBtn');
            loadingIndicator.style.display = 'block';
            submitBtn.disabled = true;
            
            // إرسال النموذج
            const formData = new FormData(this);
            
            fetch('/', {
                method: 'POST',
                body: formData
            })
            .then(response => {
                if (!response.ok) {
                    throw new Error('Network response was not ok');
                }
                return response.blob();
            })
            .then(blob => {
                // إنشاء رابط للتحميل
                const url = window.URL.createObjectURL(blob);
                const a = document.createElement('a');
                a.style.display = 'none';
                a.href = url;
                a.download = 'table.docx';
                document.body.appendChild(a);
                a.click();
                window.URL.revokeObjectURL(url);
                
                // إخفاء مؤشر التحميل
                loadingIndicator.style.display = 'none';
                submitBtn.disabled = false;
            })
            .catch(error => {
                console.error('Error:', error);
                alert('حدث خطأ أثناء إنشاء الملف. يرجى المحاولة مرة أخرى.');
                loadingIndicator.style.display = 'none';
                submitBtn.disabled = false;
            });
        });
        
        // التهيئة الأولية
        document.addEventListener('DOMContentLoaded', function() {
            updateSuggestedCriteria();
        });
    </script>
</body>
</html>
"""

# Groupes complets
group_old = [
    "أمنه عبد اللطيف","أروى يقين طنيش","اسامه بنضو","أنس الخطيب","إسراء بنمفتاح",
    "اياد بوحريه","إياد منصور عمار","المختار عبد الواحد","بادیس دقنيش","جاهد السياري",
    "رنيم العزلوك","ريتاج الطالب","رحمة الونيسي","زينب طنيش","زينب عبد الواحد",
    "سلمان الشبلي","فادي القلعاوي","الجين الزردابي","ليان الطالBI","مؤمن بنمبارك",
    "محمد أمير الحمدي","محمد الطاهر مشيري","محمد زكرياء حلاوط","مريم الذكار",
    "ملاك عبد اللطيف","منال بوحربه","هديل بن حامد","ياسمين الحاجي","ياسمين المستيسر",
    "ياسين جويد","يقين بوروحه","يوسف الشيباني","يوسف بن يحي","يونس بوصفة"
]

group_new = [
    "احلام الغليظ","أحمد التايب","أحمد الحمزي","أيمن حلموس","إدريس القرسان",
    "إسراء المرزوقي","باديس سكيب","بتول الفيتوري","تسنيم الطالب","خليل الشلاخ",
    "رضوان عبدالستار","رمزي المقدميني","رنیم خلفه","رنیم عازق","رياض لهول",
    "سيرين العربي","شيماء المورو","عبد الرحمان الوذان","عبد الرحمان بومروة",
    "الجين زهمول","محمد الطاهر بوطالب","محمد جاسم العطوي","محمد ياسين الجليدي",
    "مريم الذكار","مريم حسين","میار حسن","ميس بنصميده","ميار دباغي",
    "نزار عکار","نضال ابن غنيه","نادين مراحي","همام الغرياني","أميمة ذكار"
]

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        classe = request.form.get("classe", "سنة رابعة")
        matiere = request.form.get("matiere")
        other_subject_name = request.form.get("other_subject_name", "")
        
        # استخدام اسم المادة إذا كان "أخرى"
        if matiere == "أخرى" and other_subject_name:
            matiere = other_subject_name
        
        # Récupération des données
        criteria_json = request.form.get("criteria", "[]")
        indicators_config_json = request.form.get("indicators_config", "{}")
        
        criteria = json.loads(criteria_json)
        indicators_config = json.loads(indicators_config_json)
        
        if not criteria:
            criteria = ["معيار 1", "معيار 2", "معيار 3"]

        group_choice = request.form.get("group_choice")
        names = group_new if group_choice == "2" else group_old

        # Création du document
        doc = Document()
        
        # Configuration de la page
        section = doc.sections[0]
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.left_margin = Cm(0.8)
        section.right_margin = Cm(0.8)
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.2)
        
        # Titre principal
        title = doc.add_heading(f"جدول التقييم - {matiere}", level=1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_run = title.runs[0]
        title_run.font.size = Pt(14)
        title_run.font.bold = True
        title_run.font.name = 'Arial'

        # Sous-titre
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        subtitle_run = subtitle.add_run(f"القسم: {classe}")
        subtitle_run.font.size = Pt(11)
        subtitle_run.font.name = 'Arial'

        doc.add_paragraph().add_run().add_break()

        # حساب عدد الأعمدة الإجمالي
        total_cols = 1  # عمود الأسماء
        for criterion in criteria:
            config = indicators_config.get(criterion, {})
            if config.get('useIndicators', False):
                total_cols += 3  # 3 مؤشرات
            else:
                total_cols += 1  # معيار واحد

        # تحديد عدد صفوف الرأس
        header_rows = 2 if any(config.get('useIndicators', False) for config in indicators_config.values()) else 1
        
        # إنشاء الجدول
        table = doc.add_table(rows=header_rows, cols=total_cols)
        table.style = 'Table Grid'
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # بناء رأس الجدول
        col_index = 1  # نبدأ من العمود الثاني (بعد عمود الأسماء)
        
        # الصف الأول من الرأس
        table.rows[0].cells[0].text = "اسم التلميذ"
        table.rows[0].cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        for criterion in criteria:
            config = indicators_config.get(criterion, {})
            use_indicators = config.get('useIndicators', False)
            
            if use_indicators:
                # دمج 3 خانات للمعيار
                if col_index + 2 < total_cols:
                    table.rows[0].cells[col_index].merge(table.rows[0].cells[col_index + 2])
                
                table.rows[0].cells[col_index].text = criterion
                table.rows[0].cells[col_index].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                col_index += 3
            else:
                # معيار بدون مؤشرات - دمج الصفين
                if header_rows == 2:
                    table.rows[0].cells[col_index].merge(table.rows[1].cells[col_index])
                
                table.rows[0].cells[col_index].text = criterion
                table.rows[0].cells[col_index].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                col_index += 1

        # الصف الثاني من الرأس (للمؤشرات فقط)
        if header_rows == 2:
            col_index = 1
            for criterion in criteria:
                config = indicators_config.get(criterion, {})
                use_indicators = config.get('useIndicators', False)
                
                if use_indicators:
                    indicator_names = config.get('indicatorNames', ["مؤشر 1", "مؤشر 2", "مؤشر 3"])
                    for i in range(3):
                        table.rows[1].cells[col_index + i].text = indicator_names[i] if i < len(indicator_names) else f"مؤشر {i+1}"
                        table.rows[1].cells[col_index + i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    col_index += 3
                else:
                    col_index += 1

        # إضافة صفوف التلاميذ
        for name in names:
            row_cells = table.add_row().cells
            
            # عمود الأسماء - بدون تقطيع للسطر
            row_cells[0].text = name
            row_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            
            # تعطيل التقاطع التلقائي للنص
            for paragraph in row_cells[0].paragraphs:
                paragraph.paragraph_format.keep_together = True
                paragraph.paragraph_format.keep_with_next = False
                paragraph.paragraph_format.widow_control = False
            
            # ملء الخلايا الفارغة
            col_index = 1
            for criterion in criteria:
                config = indicators_config.get(criterion, {})
                use_indicators = config.get('useIndicators', False)
                
                if use_indicators:
                    for i in range(3):
                        row_cells[col_index + i].text = ""
                        row_cells[col_index + i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    col_index += 3
                else:
                    row_cells[col_index].text = ""
                    row_cells[col_index].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    col_index += 1

        # تطبيق التنسيق على جميع الخلايا
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.0
                    for run in paragraph.runs:
                        run.font.size = Pt(8)
                        run.font.name = 'Arial'

        # جعل الرأس عريض
        for i in range(header_rows):
            for cell in table.rows[i].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(9)

        # حساب العرض الأمثل للأعمدة
        max_name_length = max(len(name) for name in names) if names else 10
        
        # ضبط عرض الأعمدة
        col_index = 0
        for column in table.columns:
            if col_index == 0:  # عمود الأسماء
                width = min(max(Cm(2.5), Cm(max_name_length * 0.3)), Cm(6))
                column.width = width
            else:
                column.width = Cm(1.8)  # أعمدة المعايير والمؤشرات
            col_index += 1

        # إعداد RTL للجدول
        tbl = table._tbl
        tblPr = tbl.tblPr
        bidi = OxmlElement('w:bidiVisual')
        tblPr.append(bidi)

        # Sauvegarde
        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        
        from datetime import datetime
        filename = f"جدول_{matiere}_{classe}_{datetime.now().strftime('%Y%m%d')}.docx"
        return send_file(
            f,
            as_attachment=True,
            download_name=filename,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    return render_template_string(HTML)

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=True)